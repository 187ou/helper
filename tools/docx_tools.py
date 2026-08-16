"""Word 文档工具（基于 python-docx，真实可用）。"""
import logging
from typing import Any

logger = logging.getLogger(__name__)


def extract_text(path: str) -> str:
    """提取 Word 文档全部文本（段落 + 表格）。"""
    try:
        from docx import Document
        doc = Document(path)
        parts = []

        # 段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # 表格
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        result = "\n".join(parts)
        logger.info("提取 Word 文本: %s (%d 段落, %d 字符)", path, len(parts), len(result))
        return result
    except Exception as e:
        logger.error("提取 Word 文本失败 %s: %s", path, e)
        return ""


def extract_paragraphs(path: str) -> list[str]:
    """提取 Word 段落列表（保留结构）。"""
    try:
        from docx import Document
        doc = Document(path)
        return [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    except Exception as e:
        logger.error("提取 Word 段落失败 %s: %s", path, e)
        return []


def extract_tables(path: str) -> list[list[list[str]]]:
    """提取 Word 表格数据。"""
    try:
        from docx import Document
        doc = Document(path)
        tables = []
        for table in doc.tables:
            data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(data)
        return tables
    except Exception as e:
        logger.error("提取 Word 表格失败 %s: %s", path, e)
        return []


def get_doc_stats(path: str) -> dict[str, Any]:
    """获取 Word 文档统计信息。"""
    try:
        from docx import Document
        doc = Document(path)
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        char_count = sum(len(p.text) for p in doc.paragraphs)
        return {
            "paragraphs": para_count,
            "tables": table_count,
            "characters": char_count,
        }
    except Exception as e:
        logger.error("获取 Word 统计失败 %s: %s", path, e)
        return {"paragraphs": 0, "tables": 0, "characters": 0}
